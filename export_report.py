from __future__ import annotations

import argparse
import json
import re
from datetime import date, datetime
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

import yaml

from tencent_translation import load_tencent_translation_config, maybe_translate_items


ITEM_FIELDS = [
    "title_en",
    "title_zh",
    "authors",
    "year",
    "publication_date",
    "journal_or_source",
    "doi",
    "url",
    "abstract_en",
    "abstract_zh",
    "previously_seen",
    "matched_keywords",
    "relevance_score",
    "reason",
]

TABLE_COLUMNS = [
    "序号",
    "首次检索日期",
    "论文发表日期",
    "英文题目",
    "中文题目",
    "期刊 / 来源",
    "DOI",
    "链接",
    "关键词命中",
    "相关性",
    "是否以前出现过",
]


def load_config(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as file:
        return yaml.safe_load(file) or {}


def load_results(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as file:
        payload = json.load(file)

    if isinstance(payload, list):
        return {"report_date": date.today().isoformat(), "items": payload}
    return payload


def save_results(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def same_path(left: Path, right: Path) -> bool:
    return left.resolve() == right.resolve()


def save_translated_results(
    input_path: Path,
    config: dict[str, Any],
    payload: dict[str, Any],
    items: list[dict[str, Any]],
) -> None:
    payload["items"] = items
    report_date = text_value(payload.get("report_date", "")).strip()
    paths = config.get("paths", {})
    data_dir = Path(paths.get("data_dir", "data"))
    latest_path = data_dir / paths.get("latest_results_file", "latest_results.json")

    output_paths = [input_path]
    if report_date:
        dated_path = data_dir / f"{report_date}_results.json"
        if same_path(input_path, latest_path) and dated_path.exists():
            output_paths.append(dated_path)
        elif same_path(input_path, dated_path) and latest_path.exists():
            output_paths.append(latest_path)

    seen_paths: set[str] = set()
    for path in output_paths:
        key = str(path.resolve()).lower()
        if key in seen_paths:
            continue
        save_results(path, payload)
        seen_paths.add(key)


def ensure_reports_dir(config: dict[str, Any]) -> Path:
    reports_dir = Path(config.get("paths", {}).get("reports_dir", "reports"))
    reports_dir.mkdir(parents=True, exist_ok=True)
    return reports_dir


def cumulative_results_path(config: dict[str, Any]) -> Path:
    paths = config.get("paths", {})
    data_dir = Path(paths.get("data_dir", "data"))
    return data_dir / paths.get("cumulative_results_file", "cumulative_results.json")


def report_file_path(config: dict[str, Any], reports_dir: Path, key: str, default_name: str) -> Path:
    return reports_dir / config.get("paths", {}).get(key, default_name)


def text_value(value: Any) -> str:
    if isinstance(value, list):
        return "; ".join(str(item) for item in value)
    if value is None:
        return ""
    return str(value)


def bool_value(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    text = text_value(value).strip().lower()
    return text in {"1", "true", "yes", "y", "是", "previously_seen"}


def seen_text(item: dict[str, Any]) -> str:
    return "是" if bool_value(item.get("previously_seen", False)) else "否"


def parse_date_value(value: Any) -> date | None:
    text = text_value(value).strip()
    if not text:
        return None

    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00")).date()
    except ValueError:
        pass

    if re.fullmatch(r"\d{4}", text):
        return date(int(text), 1, 1)
    return None


def item_sort_date(item: dict[str, Any]) -> date:
    for field in ("first_seen_date", "retrieved_date", "last_seen_date"):
        parsed = parse_date_value(item.get(field))
        if parsed is not None:
            return parsed
    return date.min


def item_relevance_score(item: dict[str, Any]) -> float:
    try:
        return float(item.get("relevance_score", 0))
    except (TypeError, ValueError):
        return 0.0


def publication_label(item: dict[str, Any]) -> str:
    return text_value(item.get("publication_date", "")).strip() or text_value(item.get("year", "")).strip()


def retrieval_label(item: dict[str, Any]) -> str:
    for field in ("first_seen_date", "retrieved_date", "last_seen_date"):
        value = text_value(item.get(field, "")).strip()
        if value:
            return value
    return ""


def item_publication_sort_date(item: dict[str, Any]) -> date:
    return parse_date_value(item.get("publication_date")) or parse_date_value(item.get("year")) or date.min


def normalize_doi_key(value: Any) -> str:
    doi = text_value(value).strip().lower()
    doi = re.sub(r"^https?://(dx\.)?doi\.org/", "", doi)
    doi = doi.removeprefix("doi:")
    return doi.strip()


def normalize_title_key(value: Any) -> str:
    title = re.sub(r"\s+", " ", text_value(value)).strip().lower()
    return re.sub(r"[^a-z0-9]+", "", title)


def history_key(item: dict[str, Any]) -> str:
    doi = normalize_doi_key(item.get("doi"))
    if doi:
        return f"doi:{doi}"

    title = normalize_title_key(item.get("title_en") or item.get("title_zh"))
    return f"title:{title}" if title else ""


def sort_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return sorted(
        items,
        key=lambda item: (
            item_sort_date(item),
            item_relevance_score(item),
            item_publication_sort_date(item),
        ),
        reverse=True,
    )


def normalize_item(item: dict[str, Any]) -> dict[str, Any]:
    normalized = {field: text_value(item.get(field, "")) for field in ITEM_FIELDS}
    normalized["title_en"] = normalized["title_en"] or text_value(item.get("title", ""))
    normalized["abstract_en"] = normalized["abstract_en"] or text_value(item.get("abstract", ""))
    normalized["previously_seen"] = item.get("previously_seen", False)
    normalized["retrieved_date"] = text_value(item.get("retrieved_date", "") or item.get("date", ""))
    normalized["first_seen_date"] = text_value(item.get("first_seen_date", ""))
    normalized["last_seen_date"] = text_value(item.get("last_seen_date", "") or normalized["retrieved_date"])
    return normalized


def load_cumulative_results(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"items": []}
    return load_results(path)


def merge_cumulative_results(
    existing_payload: dict[str, Any],
    current_payload: dict[str, Any],
    current_items: list[dict[str, Any]],
    report_date: str,
) -> dict[str, Any]:
    merged: dict[str, dict[str, Any]] = {}

    for index, item in enumerate(existing_payload.get("items", [])):
        normalized = normalize_item(item)
        key = history_key(normalized) or f"existing:{index}"
        merged[key] = normalized

    for index, item in enumerate(current_items):
        current = normalize_item(item)
        current["retrieved_date"] = report_date
        current["last_seen_date"] = report_date
        key = history_key(current) or f"current:{report_date}:{index}"
        previous = merged.get(key)

        if previous is not None:
            current["first_seen_date"] = previous.get("first_seen_date") or previous.get("retrieved_date") or report_date
            current["previously_seen"] = True
        else:
            current["first_seen_date"] = report_date

        merged[key] = current

    items = sort_items(list(merged.values()))
    return {
        "generated_at": current_payload.get("generated_at", ""),
        "latest_report_date": report_date,
        "lookback_days": current_payload.get("lookback_days", ""),
        "cutoff_date": current_payload.get("cutoff_date", ""),
        "sources": current_payload.get("sources", {}),
        "items": items,
    }


def unique_doi_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen: set[str] = set()
    doi_items: list[dict[str, Any]] = []
    for item in items:
        doi = text_value(item.get("doi", "")).strip()
        key = doi.lower()
        if not doi or key in seen:
            continue
        seen.add(key)
        doi_items.append(item)
    return doi_items


def format_source_counts(sources: dict[str, Any]) -> str:
    parts = []
    for source, count in sorted(sources.items()):
        parts.append(f"{source}={count}")
    return ", ".join(parts) if parts else "No source counts available"


def item_text_blob(item: dict[str, Any]) -> str:
    parts = [
        item.get("title_en", ""),
        item.get("abstract_en", ""),
        item.get("matched_keywords", ""),
        item.get("reason", ""),
    ]
    return " ".join(text_value(part) for part in parts).lower()


def has_any(text: str, needles: list[str]) -> bool:
    return any(needle.lower() in text for needle in needles)


def relevance_reasons(item: dict[str, Any]) -> list[str]:
    text = item_text_blob(item)
    organic_terms = [
        "organic molecule",
        "dye",
        "chromophore",
        "photosensitizer",
        "molecular antenna",
        "organic antenna",
        "dye-sensitized",
    ]
    gold_terms = [
        "gold nanoparticle",
        "gold nanoparticles",
        "au nanoparticle",
        "au nanoparticles",
        "au nanorod",
        "gold nanorod",
        "gold nanostar",
        "gold nanoshell",
        "gold shell",
        "gold film",
        "gold nanoarray",
        "plasmonic gold",
        "gold nanostructure",
    ]
    plasmon_terms = [
        "plasmon",
        "lspr",
        "localized surface plasmon resonance",
        "plasmon-mediated energy transfer",
        "plasmon-enhanced luminescence",
        "plasmonic enhancement",
    ]
    lanthanide_terms = [
        "lanthanide",
        "rare-earth",
        "upconversion nanoparticle",
        "upconversion nanoparticles",
        "ucnp",
        "naerf4",
        "nayf4",
        "er3+",
        "er 3+",
        "erbium",
    ]
    reasons: list[str] = []

    if has_any(text, organic_terms) and has_any(text, gold_terms):
        reasons.append("包含 organic molecule / dye 与 gold nanoparticle coupling 相关关键词")
    if "plasmon-mediated energy transfer" in text or (has_any(text, plasmon_terms) and "energy transfer" in text):
        reasons.append("包含 plasmon-mediated energy transfer 或 plasmonic energy-transfer 线索")
    if has_any(text, plasmon_terms) and has_any(text, lanthanide_terms):
        reasons.append("包含 plasmon-enhanced lanthanide luminescence / upconversion 相关线索")
    if has_any(text, gold_terms):
        reasons.append("包含 gold nanorod / Au nanoparticle / plasmonic gold nanostructure 相关线索")
    if has_any(text, organic_terms) and has_any(text, lanthanide_terms):
        reasons.append("包含 organic antenna sensitization of lanthanide nanoparticles 线索")
    if "molecule-to-metal energy transfer" in text:
        reasons.append("包含 molecule-to-metal energy transfer")
    if "metal-to-lanthanide energy transfer" in text:
        reasons.append("包含 metal-to-lanthanide energy transfer")
    if "plasmon-exciton coupling" in text:
        reasons.append("包含 plasmon-exciton coupling")
    if "4f" in text or "energy level" in text:
        reasons.append("包含 Er3+ specific energy level / 4f transition / lanthanide energy-transfer 线索")

    if not reasons:
        matched = text_value(item.get("matched_keywords", "")).strip()
        if matched:
            reasons.append(f"关键词命中：{matched}")
        else:
            existing_reason = text_value(item.get("reason", "")).strip()
            reasons.append(existing_reason or "根据题目或摘要命中的宽关键词列为候选")

    return reasons[:6]


def _rgb_color(hex_value: str) -> Any:
    from docx.shared import RGBColor

    clean = hex_value.strip().lstrip("#")
    return RGBColor(int(clean[0:2], 16), int(clean[2:4], 16), int(clean[4:6], 16))


def _set_rfonts(element: Any, latin_font: str = "Times New Roman", east_asia_font: str = "SimSun") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def _set_run_font(
    run: Any,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    latin_font: str = "Times New Roman",
    east_asia_font: str = "SimSun",
) -> None:
    from docx.shared import Pt

    run.font.name = latin_font
    _set_rfonts(run._element, latin_font, east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _rgb_color(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(
    style: Any,
    *,
    size: float,
    color: str,
    bold: bool | None = None,
    latin_font: str = "Times New Roman",
    east_asia_font: str = "SimSun",
) -> None:
    from docx.shared import Pt

    style.font.name = latin_font
    _set_rfonts(style._element, latin_font, east_asia_font)
    style.font.size = Pt(size)
    style.font.color.rgb = _rgb_color(color)
    if bold is not None:
        style.font.bold = bold


def _clear_paragraph(paragraph: Any) -> None:
    paragraph._p.clear_content()


def _set_paragraph_bottom_border(paragraph: Any, *, color: str = "D7DEE8", size: str = "12", space: str = "10") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def _set_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        margin = tc_mar.find(qn(f"w:{edge}"))
        if margin is None:
            margin = OxmlElement(f"w:{edge}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _set_cell_borders(cell: Any, *, color: str = "D7DEE8", size: str = "4") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end"):
        border = tc_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tc_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _set_table_geometry(table: Any, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for child in list(tbl_pr):
        if child.tag in {qn("w:tblW"), qn("w:tblInd"), qn("w:tblLayout")}:
            tbl_pr.remove(child)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    for row in table.rows:
        for column_index, width in enumerate(widths_dxa):
            cell = row.cells[column_index]
            tc_pr = cell._tc.get_or_add_tcPr()
            for child in list(tc_pr):
                if child.tag == qn("w:tcW"):
                    tc_pr.remove(child)
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)


def _repeat_table_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _style_cell_text(
    cell: Any,
    *,
    size: float = 9,
    color: str = "1F2937",
    bold: bool = False,
    align: Any = None,
    line_spacing: float = 1.1,
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    _set_cell_borders(cell)
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = line_spacing
        for run in paragraph.runs:
            _set_run_font(run, size=size, color=color, bold=bold)


def _add_page_number(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    _set_run_font(run, size=9, color="6B7280")


def _configure_section(section: Any, *, landscape: bool = False) -> None:
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches

    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Inches(11 if landscape else 8.5)
    section.page_height = Inches(8.5 if landscape else 11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def _configure_document_styles(document: Any) -> None:
    from docx.shared import Inches, Pt

    styles = document.styles

    normal = styles["Normal"]
    _set_style_font(normal, size=11, color="1F2937")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    _set_style_font(title, size=24, color="0B2545", bold=True)
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.line_spacing = 1.05

    heading_1 = styles["Heading 1"]
    _set_style_font(heading_1, size=16, color="2E74B5", bold=True)
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.line_spacing = 1.25
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    _set_style_font(heading_2, size=13, color="2E74B5", bold=True)
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(7)
    heading_2.paragraph_format.line_spacing = 1.25
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = styles["Heading 3"]
    _set_style_font(heading_3, size=12, color="1F4D78", bold=True)
    heading_3.paragraph_format.space_before = Pt(10)
    heading_3.paragraph_format.space_after = Pt(5)
    heading_3.paragraph_format.line_spacing = 1.25
    heading_3.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    _set_style_font(bullet, size=10.5, color="1F2937")
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25


def _set_running_header_footer(section: Any, report_date: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    _clear_paragraph(header_paragraph)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run(f"累计文献检索报告 · 更新至 {report_date}")
    _set_run_font(header_run, size=9, color="6B7280", bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    _clear_paragraph(footer_paragraph)
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_run = footer_paragraph.add_run("第 ")
    _set_run_font(footer_run, size=9, color="6B7280")
    _add_page_number(footer_paragraph)
    end_run = footer_paragraph.add_run(" 页")
    _set_run_font(end_run, size=9, color="6B7280")


def _add_title_block(document: Any, report_date: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    run = paragraph.add_run(f"累计文献检索报告 - 更新至 {report_date}")
    _set_run_font(run, size=24, color="0B2545", bold=True)
    _set_paragraph_bottom_border(paragraph, color="B9CBE0", size="14", space="12")


def _add_report_overview(document: Any, payload: dict[str, Any], items: list[dict[str, Any]]) -> None:
    generated_at = text_value(payload.get("generated_at", ""))
    lookback_days = text_value(payload.get("lookback_days", "")).strip()
    cutoff_date = text_value(payload.get("cutoff_date", "")).strip()
    sources = payload.get("sources", {})

    search_window = ""
    if lookback_days or cutoff_date:
        search_window = f"最近 {lookback_days or '?'} 天"
        if cutoff_date:
            search_window += f"，起始日期 {cutoff_date}"

    _add_section_heading(document, "检索概览")
    _add_metadata_table(
        document,
        [
            ("生成时间", generated_at),
            ("检索窗口", search_window),
            ("候选文献数", len(items)),
            ("来源抓取数", format_source_counts(sources if isinstance(sources, dict) else {})),
        ],
    )


def _add_section_heading(document: Any, text: str) -> None:
    document.add_heading(text, level=1)


def _add_doi_list(document: Any, items: list[dict[str, Any]]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doi_items = unique_doi_items(items)
    if not doi_items:
        paragraph = document.add_paragraph("No DOI available")
        for run in paragraph.runs:
            _set_run_font(run)
        return

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    column_widths = [1300, 1300, 6760]
    _set_table_geometry(table, column_widths)
    _repeat_table_header(table.rows[0])
    for column_index, column_name in enumerate(["首次检索日期", "论文发表日期", "DOI"]):
        cell = table.rows[0].cells[column_index]
        cell.text = column_name
        _set_cell_shading(cell, "2E74B5")
        _style_cell_text(cell, size=9, color="FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.05)

    for row_index, item in enumerate(doi_items, start=1):
        row = table.add_row()
        fill = "F8FAFC" if row_index % 2 == 0 else "FFFFFF"
        values = [
            retrieval_label(item),
            publication_label(item),
            text_value(item.get("doi", "")).strip(),
        ]
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.text = value
            _set_cell_shading(cell, fill)
            align = WD_ALIGN_PARAGRAPH.LEFT if column_index == 2 else WD_ALIGN_PARAGRAPH.CENTER
            _style_cell_text(cell, size=9, color="0B2545", align=align, line_spacing=1.15)
    _set_table_geometry(table, column_widths)


def _add_summary_table(document: Any, items: list[dict[str, Any]]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = document.add_table(rows=1, cols=len(TABLE_COLUMNS))
    table.style = "Table Grid"
    column_widths = [450, 850, 850, 2250, 1850, 1150, 1300, 1450, 850, 600, 600]
    _set_table_geometry(table, column_widths)
    _repeat_table_header(table.rows[0])

    for column_index, column_name in enumerate(TABLE_COLUMNS):
        cell = table.rows[0].cells[column_index]
        cell.text = column_name
        _set_cell_shading(cell, "2E74B5")
        _style_cell_text(cell, size=8, color="FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.05)

    for row_index, item in enumerate(items, start=1):
        row = table.add_row()
        fill = "F8FAFC" if row_index % 2 == 0 else "FFFFFF"
        values = [
            str(row_index),
            retrieval_label(item),
            publication_label(item),
            text_value(item.get("title_en", "")),
            text_value(item.get("title_zh", "")),
            text_value(item.get("journal_or_source", "")),
            text_value(item.get("doi", "")),
            text_value(item.get("url", "")),
            text_value(item.get("matched_keywords", "")),
            text_value(item.get("relevance_score", "")),
            seen_text(item),
        ]
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.text = value
            _set_cell_shading(cell, fill)
            align = WD_ALIGN_PARAGRAPH.CENTER if column_index in {0, 1, 2, 9, 10} else WD_ALIGN_PARAGRAPH.LEFT
            _style_cell_text(cell, size=7.5, color="1F2937", align=align, line_spacing=1.05)
    _set_table_geometry(table, column_widths)


def _add_metadata_table(document: Any, rows: list[tuple[str, Any]]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table, [1700, 7660])
    for row_index, (label, value) in enumerate(rows):
        label_cell = table.rows[row_index].cells[0]
        value_cell = table.rows[row_index].cells[1]
        label_cell.text = f"{label}："
        value_cell.text = text_value(value)
        _set_cell_shading(label_cell, "E8EEF5")
        _set_cell_shading(value_cell, "FFFFFF")
        _style_cell_text(label_cell, size=9, color="1F4D78", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _style_cell_text(value_cell, size=9, color="1F2937", align=WD_ALIGN_PARAGRAPH.LEFT)


def _add_label_paragraph(document: Any, text: str) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    _set_run_font(run, size=10.5, color="1F4D78", bold=True)


def _add_body_paragraph(document: Any, text: Any) -> None:
    paragraph = document.add_paragraph(text_value(text))
    for run in paragraph.runs:
        _set_run_font(run, size=10.5, color="1F2937")


def _clean_word_font_theme(path: Path) -> None:
    temp_path = path.with_suffix(path.suffix + ".tmp")
    replacements = {
        "Calibri": "Times New Roman",
        "Microsoft YaHei": "SimSun",
    }
    target_parts = {"word/fontTable.xml", "word/theme/theme1.xml", "word/styles.xml"}

    with ZipFile(path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename in target_parts:
                text = data.decode("utf-8")
                for old, new in replacements.items():
                    text = text.replace(old, new)
                data = text.encode("utf-8")
            target.writestr(info, data)

    temp_path.replace(path)


def write_word(path: Path, payload: dict[str, Any], items: list[dict[str, Any]]) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
    except ImportError as error:
        raise RuntimeError("python-docx is required to generate Word reports. Run: pip install -r requirements.txt") from error

    report_date = text_value(payload.get("latest_report_date") or payload.get("report_date") or date.today().isoformat())
    document = Document()
    _configure_section(document.sections[0], landscape=False)
    _configure_document_styles(document)
    _set_running_header_footer(document.sections[0], report_date)
    _add_title_block(document, report_date)

    _add_report_overview(document, payload, items)

    _add_section_heading(document, "一、累计 DOI 清单")
    _add_doi_list(document, items)

    if items:
        summary_section = document.add_section(WD_SECTION.NEW_PAGE)
        _configure_section(summary_section, landscape=True)
        _set_running_header_footer(summary_section, report_date)

    _add_section_heading(document, "二、候选文献总表")
    if items:
        _add_summary_table(document, items)
    else:
        document.add_paragraph("No candidate literature was found.")

    if items:
        detail_section = document.add_section(WD_SECTION.NEW_PAGE)
        _configure_section(detail_section, landscape=False)
        _set_running_header_footer(detail_section, report_date)

    _add_section_heading(document, "三、详细文献信息")
    if not items:
        document.add_paragraph("No candidate literature was found.")

    for index, item in enumerate(items, start=1):
        title_en = text_value(item.get("title_en", "")).strip() or "(No English title)"
        document.add_heading(f"{index}. {title_en}", level=2)
        _add_metadata_table(
            document,
            [
                ("中文题目", item.get("title_zh", "")),
                ("DOI", item.get("doi", "")),
                ("链接", item.get("url", "")),
                ("首次检索日期", retrieval_label(item)),
                ("论文发表日期", publication_label(item)),
                ("期刊 / 来源", item.get("journal_or_source", "")),
                ("作者", item.get("authors", "")),
                ("最近检索日期", item.get("last_seen_date", "")),
                ("相关性分数", item.get("relevance_score", "")),
                ("关键词命中", item.get("matched_keywords", "")),
                ("是否以前出现过", seen_text(item)),
            ],
        )

        _add_label_paragraph(document, "英文摘要：")
        _add_body_paragraph(document, item.get("abstract_en", ""))
        _add_label_paragraph(document, "中文摘要：")
        _add_body_paragraph(document, item.get("abstract_zh", ""))
        _add_label_paragraph(document, "可能相关原因：")
        for reason in relevance_reasons(item):
            paragraph = document.add_paragraph(reason, style="List Bullet")
            for run in paragraph.runs:
                _set_run_font(run, size=10.5, color="1F2937")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    _clean_word_font_theme(path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export literature search results to a Word report.")
    parser.add_argument("--config", default="config.yaml", help="Path to config.yaml.")
    parser.add_argument("--input", default=None, help="Input JSON file. Defaults to data/latest_results.json.")
    parser.add_argument("--date", default=None, help="Override report date in YYYY-MM-DD format.")
    parser.add_argument("--no-translate", action="store_true", help="Do not use optional Tencent Cloud translation even if .env has Tencent Cloud credentials.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    config = load_config(Path(args.config))
    paths = config.get("paths", {})
    input_path = Path(args.input or Path(paths.get("data_dir", "data")) / paths.get("latest_results_file", "latest_results.json"))
    payload = load_results(input_path)

    report_date = args.date or payload.get("report_date") or date.today().isoformat()
    payload["report_date"] = report_date
    items = [normalize_item(item) for item in payload.get("items", [])]
    translation_cache_path = Path(paths.get("data_dir", "data")) / "translation_cache.json"
    translation_config = None if args.no_translate else load_tencent_translation_config(cache_path=translation_cache_path)
    items = maybe_translate_items(items, translation_config)
    items = sort_items(items)
    save_translated_results(input_path, config, payload, items)

    reports_dir = ensure_reports_dir(config)
    history_path = cumulative_results_path(config)
    cumulative_payload = merge_cumulative_results(load_cumulative_results(history_path), payload, items, report_date)
    save_results(history_path, cumulative_payload)

    cumulative_items = cumulative_payload["items"]
    word_path = report_file_path(config, reports_dir, "word_report_file", "literature_report.docx")

    write_word(word_path, cumulative_payload, cumulative_items)

    print(f"Saved Word report to {word_path}")
    print(f"Updated cumulative results at {history_path}")


if __name__ == "__main__":
    main()
